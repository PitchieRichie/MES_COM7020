#!/usr/bin/env python3
"""
MES assignment submission generator (single entry point).

Produces:
  - support/output/figs/fig1..fig6.png
  - support/output/metrics.csv
  - updates Designing a Cloud - MES_014517.docx (figures + academic formatting)

Layout:
  _temp/                          <- report, dataset, notebook
  _temp/support/                  <- this script, outputs, backups

Usage:
  python support/generate_mes_submission.py
"""

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from sklearn.ensemble import GradientBoostingRegressor
from sklearn.metrics import mean_absolute_error, mean_absolute_percentage_error, r2_score

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SUPPORT = Path(__file__).resolve().parent
ROOT = SUPPORT.parent
DOCX = ROOT / "Designing a Cloud - MES_014517.docx"
FIG_DIR = SUPPORT / "output" / "figs"
FIG_DIR.mkdir(parents=True, exist_ok=True)

# Warm red limited set (shared across all figures)
INK = "#3B1A16"
PRIMARY = "#C45C4A"
SOFT = "#F7EBE7"
PAPER = "#FBF7F5"
LINE = "#E2D4CF"
MUTE = "#7A5E58"
WHITE = "#FFFFFF"
ACCENT = "#E8A07A"

# Back-compat aliases used in older call sites / notebook imports
TEAL = PRIMARY
NAVY2 = "#8B3A32"

plt.rcParams.update(
    {
        "figure.dpi": 150,
        "savefig.dpi": 200,
        "font.family": "sans-serif",
        "font.sans-serif": ["Helvetica Neue", "Arial", "DejaVu Sans"],
        "axes.unicode_minus": False,
    }
)


# ---------------------------------------------------------------------------
# Shared drawing helpers — keep every diagram on the same visual system
# ---------------------------------------------------------------------------
def canvas(size):
    fig, ax = plt.subplots(figsize=size, facecolor=PAPER)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.set_facecolor(PAPER)
    fig.subplots_adjust(left=0.03, right=0.97, top=0.95, bottom=0.04)
    return fig, ax


def txt(ax, x, y, s, size=8.5, color=INK, weight="bold", ha="center", va="center"):
    # Map soft weights to fonts available on macOS Helvetica Neue
    fw = {"medium": "normal", "semibold": "bold"}.get(weight, weight)
    ax.text(
        x,
        y,
        s,
        transform=ax.transAxes,
        fontsize=size,
        color=color,
        fontweight=fw,
        ha=ha,
        va=va,
        zorder=5,
        clip_on=False,
    )


def card(ax, x, y, w, h, title, subtitle=None, face=WHITE, edge=PRIMARY, lw=1.35, title_size=8.0, sub_size=6.4, lines=None):
    """Rounded card; title/subtitle (or extra lines) centred inside the box."""
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.004,rounding_size=0.012",
            facecolor=face,
            edgecolor=edge,
            linewidth=lw,
            transform=ax.transAxes,
            clip_on=False,
            zorder=2,
        )
    )
    cx, cy = x + w / 2.0, y + h / 2.0
    extras = list(lines or [])
    if subtitle:
        extras = [subtitle] + extras
    if extras:
        # Title above centre; supporting lines stacked below
        block = 1 + len(extras)
        top = cy + (block - 1) * 0.016
        txt(ax, cx, top, title, size=title_size, weight="bold")
        for i, line in enumerate(extras):
            txt(ax, cx, top - (i + 1) * 0.032, line, size=sub_size, color=MUTE, weight="medium")
    else:
        txt(ax, cx, cy, title, size=title_size, weight="bold")
    return {"x": x, "y": y, "w": w, "h": h, "cx": cx, "cy": cy, "r": x + w, "t": y + h}


def arrow(ax, x0, y0, x1, y1, color=MUTE, lw=1.3):
    ax.add_patch(
        FancyArrowPatch(
            (x0, y0),
            (x1, y1),
            transform=ax.transAxes,
            arrowstyle="-|>",
            mutation_scale=9,
            lw=lw,
            color=color,
            shrinkA=0,
            shrinkB=0,
            zorder=3,
            clip_on=False,
        )
    )


def h_link(ax, a, b, color=MUTE, y=None):
    yy = a["cy"] if y is None else y
    arrow(ax, a["r"] + 0.005, yy, b["x"] - 0.005, yy, color=color)


def v_link(ax, a, b, color=MUTE, x=None):
    xx = a["cx"] if x is None else x
    arrow(ax, xx, a["y"] - 0.005, xx, b["t"] + 0.005, color=color)


def style_chart(ax, title: str):
    ax.set_facecolor(PAPER)
    ax.set_title(title, color=INK, fontsize=11, pad=10, fontweight="bold")
    ax.tick_params(colors=MUTE, labelsize=8)
    ax.xaxis.label.set_color(MUTE)
    ax.yaxis.label.set_color(MUTE)
    ax.grid(alpha=0.28, color=LINE)
    for spine in ax.spines.values():
        spine.set_color(LINE)
    handles, labels = ax.get_legend_handles_labels()
    if handles:
        leg = ax.legend(frameon=False, fontsize=8)
        for t in leg.get_texts():
            t.set_color(INK)


def save(fig, name: str):
    path = FIG_DIR / name
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor=PAPER, pad_inches=0.18)
    plt.close(fig)
    print("Saved", path.name)
    return path


# ---------------------------------------------------------------------------
# Figure 1 — Shared responsibility matrix
# ---------------------------------------------------------------------------
def fig1_shared_responsibility():
    layers = [
        "Application & data",
        "Runtime & middleware",
        "Operating system",
        "Virtualisation",
        "Servers / storage / network",
        "Physical facility",
    ]
    matrix = {
        "On-premises": [False] * 6,
        "IaaS": [False, False, False, True, True, True],
        "PaaS": [False, True, True, True, True, True],
        "SaaS": [True] * 6,
    }
    models = list(matrix.keys())

    fig, ax = canvas((11.0, 5.6))
    txt(ax, 0.5, 0.96, "Figure 1. Shared responsibility across cloud service models", size=12, weight="bold")

    label_right = 0.27
    left, top = 0.30, 0.82
    col_w, row_h = 0.150, 0.090
    gap_x, gap_y = 0.022, 0.016

    for j, m in enumerate(models):
        cx = left + j * (col_w + gap_x) + col_w / 2
        txt(ax, cx, top + 0.045, m, size=9.2, weight="bold")

    for i, layer in enumerate(layers):
        y = top - (i + 1) * row_h - i * gap_y
        txt(ax, label_right - 0.015, y + row_h / 2, layer, size=8, ha="right", color=MUTE, weight="medium")
        for j, m in enumerate(models):
            x = left + j * (col_w + gap_x)
            provider = matrix[m][i]
            face = PRIMARY if provider else ACCENT
            ax.add_patch(
                FancyBboxPatch(
                    (x, y),
                    col_w,
                    row_h,
                    boxstyle="round,pad=0.002,rounding_size=0.01",
                    facecolor=face,
                    edgecolor=face,
                    linewidth=0,
                    transform=ax.transAxes,
                    clip_on=False,
                    zorder=2,
                )
            )
            txt(
                ax,
                x + col_w / 2,
                y + row_h / 2,
                "Provider" if provider else "MES",
                size=7.3,
                color=WHITE,
                weight="bold",
            )

    legend_y = top - len(layers) * (row_h + gap_y) - 0.04
    sw = sh = 0.026
    items = [
        (left, ACCENT, "Managed by MES"),
        (left + 2 * (col_w + gap_x), PRIMARY, "Managed by cloud provider"),
    ]
    for lx, face, label in items:
        ax.add_patch(
            FancyBboxPatch(
                (lx, legend_y),
                sw,
                sh,
                boxstyle="round,pad=0.001,rounding_size=0.006",
                facecolor=face,
                edgecolor=face,
                linewidth=0,
                transform=ax.transAxes,
                zorder=2,
            )
        )
        txt(ax, lx + sw + 0.012, legend_y + sh / 2, label, size=8, ha="left", color=MUTE, weight="medium")

    save(fig, "fig1_shared_responsibility.png")


# ---------------------------------------------------------------------------
# Figure 2 — Architecture (aligned columns; every node linked)
# ---------------------------------------------------------------------------
def fig2_architecture():
    """Layered architecture on a fixed grid; connectors never cross cards."""
    fig, ax = canvas((12.2, 7.6))
    txt(ax, 0.5, 0.975, "Figure 2. Proposed cloud-enabled big data architecture for MES", size=12, weight="bold")

    cols = [0.04, 0.26, 0.48, 0.70]
    bw = 0.18
    y_top, h_top = 0.78, 0.13
    y_lake, h_lake = 0.56, 0.11
    y_serve, h_serve = 0.38, 0.10
    y_cons, h_cons = 0.22, 0.09
    y_comp, h_comp = 0.05, 0.07

    txt(ax, 0.04, 0.94, "Coral = stream path    Ink = batch path", size=7.2, color=MUTE, weight="medium", ha="left")

    sources = card(
        ax, cols[0], y_top, bw, h_top, "Source systems",
        lines=["Meters · EV · Solar", "IoT · Billing · Weather"],
        face=SOFT, edge=PRIMARY, title_size=8.0, sub_size=6.1,
    )
    h_half = (h_top - 0.012) / 2
    stream_in = card(ax, cols[1], y_top + h_half + 0.012, bw, h_half, "Stream ingest", "event broker", face=SOFT, edge=PRIMARY, title_size=7.4)
    batch_in = card(ax, cols[1], y_top, bw, h_half, "Batch ingest", "scheduled load", face=WHITE, edge=INK, title_size=7.4)
    raw = card(ax, cols[2], y_top, bw, h_top, "Raw / landing", "immutable object store", face=SOFT, edge=PRIMARY, title_size=8.0)
    stream_pr = card(ax, cols[3], y_top + h_half + 0.012, bw, h_half, "Stream process", "near-real-time", face=SOFT, edge=PRIMARY, title_size=7.4)
    batch_pr = card(ax, cols[3], y_top, bw, h_half, "Batch process", "curation jobs", face=WHITE, edge=INK, title_size=7.4)

    lake_x, lake_w = cols[1], (cols[3] + bw) - cols[1]
    lake = card(
        ax, lake_x, y_lake, lake_w, h_lake,
        "Curated lakehouse (open formats)",
        "cleansed · conformed · versioned tables",
        face=SOFT, edge=PRIMARY, lw=1.7, title_size=9.0,
    )

    serve_gap = 0.016
    serve_w = (lake_w - 2 * serve_gap) / 3.0
    serves = [
        card(ax, lake_x + i * (serve_w + serve_gap), y_serve, serve_w, h_serve, t, s, face=WHITE, edge=PRIMARY, title_size=7.2)
        for i, (t, s) in enumerate([
            ("Warehouse / marts", "reporting"),
            ("ML / forecasting", "PoC workbench"),
            ("APIs / events", "interoperability"),
        ])
    ]
    cons = [
        card(ax, lake_x + i * (serve_w + serve_gap), y_cons, serve_w, h_cons, t, None, face=SOFT, edge=PRIMARY, title_size=7.2)
        for i, t in enumerate(["Grid ops", "Billing / self-service", "Legacy / vendors"])
    ]

    gov = card(
        ax, cols[0], y_cons, bw, (y_lake + h_lake) - y_cons,
        "Governance",
        lines=["IAM · encryption", "lineage · GDPR", "catalogue · quality"],
        face=WHITE, edge=INK, title_size=8.0,
    )
    comp = card(
        ax, cols[0], y_comp, (cols[3] + bw) - cols[0], h_comp,
        "Compliance envelope: UK GDPR · DPA 2018 · NIS · Ofgem · ISO/IEC 27001",
        None, face=SOFT, edge=PRIMARY, title_size=7.1,
    )

    h_link(ax, sources, stream_in, color=PRIMARY, y=stream_in["cy"])
    h_link(ax, sources, batch_in, color=INK, y=batch_in["cy"])
    h_link(ax, stream_in, raw, color=PRIMARY, y=stream_in["cy"])
    h_link(ax, batch_in, raw, color=INK, y=batch_in["cy"])
    h_link(ax, raw, stream_pr, color=PRIMARY, y=stream_pr["cy"])
    h_link(ax, raw, batch_pr, color=INK, y=batch_pr["cy"])

    # Process exits: stream via right rail, batch straight down (distinct lake entry points)
    x_rail = min(0.97, stream_pr["r"] + 0.02)
    x_stream_in = lake["x"] + lake["w"] * 0.72
    x_batch_in = lake["x"] + lake["w"] * 0.88
    ax.plot(
        [stream_pr["r"] + 0.002, x_rail, x_rail, x_stream_in],
        [stream_pr["cy"], stream_pr["cy"], lake["t"] + 0.014, lake["t"] + 0.014],
        transform=ax.transAxes, color=PRIMARY, lw=1.25, solid_capstyle="round", zorder=3,
    )
    arrow(ax, x_stream_in, lake["t"] + 0.014, x_stream_in, lake["t"] + 0.004, color=PRIMARY, lw=1.25)
    ax.plot(
        [batch_pr["cx"], batch_pr["cx"], x_batch_in],
        [batch_pr["y"] - 0.004, lake["t"] + 0.014, lake["t"] + 0.014],
        transform=ax.transAxes, color=INK, lw=1.25, solid_capstyle="round", zorder=3,
    )
    arrow(ax, x_batch_in, lake["t"] + 0.014, x_batch_in, lake["t"] + 0.004, color=INK, lw=1.25)

    for s in serves:
        v_link(ax, lake, s, color=MUTE, x=s["cx"])
    for s, cnode in zip(serves, cons):
        v_link(ax, s, cnode, color=MUTE)
    for cnode in cons:
        v_link(ax, cnode, comp, color=MUTE, x=cnode["cx"])
    h_link(ax, gov, lake, color=INK, y=lake["cy"])
    save(fig, "fig2_architecture.png")


def fig3_dataflow():
    """Five equal columns, three equal cards each, shared spine + feedback loop."""
    fig, ax = canvas((12.4, 6.0))
    txt(ax, 0.5, 0.965, "Figure 3. End-to-end data flow from field devices to business decisions", size=11.5, weight="bold")

    stages = [
        ("Sources", [("Field devices", "Meters · EV · solar / IoT"), ("Enterprise", "Billing · CMMS"), ("External feeds", "Weather · tariffs")], SOFT, PRIMARY),
        ("Ingest", [("Stream ingest", "event broker"), ("Batch ingest", "connectors"), ("Normalise", "schema registry")], WHITE, INK),
        ("Lakehouse", [("Raw zone", "immutable landing"), ("Curated zone", "cleansed tables"), ("Gold zone", "feature / mart ready")], SOFT, PRIMARY),
        ("Analytics", [("Stream analytics", "fault alerts"), ("Batch / ML", "forecast · profiles"), ("Reporting", "billing extracts")], WHITE, INK),
        ("Decisions", [("Grid ops", "control room"), ("Customer apps", "self-service"), ("Regulatory", "audit packs")], SOFT, PRIMARY),
    ]

    n = len(stages)
    col_w = 0.150
    left, right = 0.03, 0.97
    usable = right - left - col_w
    xs = [left + i * (usable / (n - 1)) for i in range(n)]

    card_h, gap = 0.125, 0.018
    n_cards = 3
    frame_h = n_cards * card_h + (n_cards - 1) * gap
    frame_top = 0.80
    spine_y = frame_top - frame_h / 2

    col_cards = []
    for xi, (title, items, face, edge) in zip(xs, stages):
        # Column panel for grouping
        ax.add_patch(
            FancyBboxPatch(
                (xi - 0.012, frame_top - frame_h - 0.025),
                col_w + 0.024,
                frame_h + 0.055,
                boxstyle="round,pad=0.004,rounding_size=0.012",
                facecolor=WHITE if edge == INK else SOFT,
                edgecolor=LINE,
                linewidth=1.0,
                transform=ax.transAxes,
                clip_on=False,
                zorder=1,
            )
        )
        txt(ax, xi + col_w / 2, 0.875, title, size=9, weight="bold", color=MUTE)
        cards = []
        for i, (h, s) in enumerate(items):
            y = frame_top - card_h - i * (card_h + gap)
            cards.append(card(ax, xi, y, col_w, card_h, h, s, face=face, edge=edge, title_size=7.4, sub_size=6.0))
        col_cards.append(cards)

    # Spine arrows between column panels
    for i in range(n - 1):
        arrow(ax, xs[i] + col_w + 0.012, spine_y, xs[i + 1] - 0.012, spine_y, color=MUTE, lw=1.4)

    # Feedback loop under panels
    y_fb = 0.09
    first = col_cards[0][-1]
    last = col_cards[-1][-1]
    ax.plot(
        [last["cx"], last["cx"], first["cx"], first["cx"]],
        [last["y"] - 0.01, y_fb, y_fb, first["y"] - 0.01],
        transform=ax.transAxes, color=PRIMARY, lw=1.25, solid_capstyle="round", zorder=3,
    )
    arrow(ax, first["cx"], y_fb + 0.015, first["cx"], first["y"] - 0.01, color=PRIMARY, lw=1.25)
    txt(ax, 0.5, y_fb + 0.03, "Continuous pipeline monitoring and data-quality feedback", size=7.8, color=PRIMARY, weight="bold")
    save(fig, "fig3_dataflow.png")


# ---------------------------------------------------------------------------
# Figures 4-6 — PoC charts
# ---------------------------------------------------------------------------
def run_poc_and_figs_4_to_6():
    rng = np.random.default_rng(42)
    periods = 120 * 48
    idx = pd.date_range("2026-01-01", periods=periods, freq="30min")
    df = pd.DataFrame(index=idx)
    df["hour"] = df.index.hour + df.index.minute / 60
    df["dow"] = df.index.dayofweek
    df["doy"] = df.index.dayofyear
    df["is_weekend"] = (df["dow"] >= 5).astype(int)
    seasonal_temp = 8 + 6 * np.sin(2 * np.pi * (df["doy"] - 30) / 365)
    df["temperature_c"] = seasonal_temp + rng.normal(0, 1.5, periods)
    morning = 1.6 * np.exp(-((df["hour"] - 7.5) ** 2) / (2 * 1.2**2))
    evening = 2.4 * np.exp(-((df["hour"] - 18.5) ** 2) / (2 * 1.6**2))
    overnight = 0.9 + 0.15 * np.sin(2 * np.pi * df["hour"] / 24)
    weekend = np.where(df["is_weekend"] == 1, 0.85, 1.0)
    heating = np.clip(15 - df["temperature_c"], 0, None) * 0.05
    df["load_kwh"] = ((overnight + morning + evening) * weekend + heating + rng.normal(0, 0.12, periods)).clip(0.05)
    fault_idx = rng.choice(periods, size=6, replace=False)
    df.iloc[fault_idx, df.columns.get_loc("load_kwh")] *= 0.2
    df["lag_48"] = df["load_kwh"].shift(48)
    df["lag_336"] = df["load_kwh"].shift(336)
    df["roll_mean_48"] = df["load_kwh"].rolling(48).mean()
    df = df.dropna()

    features = ["hour", "dow", "is_weekend", "temperature_c", "lag_48", "lag_336", "roll_mean_48"]
    target = "load_kwh"
    split = int(len(df) * 0.85)
    train, test = df.iloc[:split], df.iloc[split:]
    model = GradientBoostingRegressor(n_estimators=250, max_depth=3, learning_rate=0.05, random_state=42)
    model.fit(train[features], train[target])
    pred = model.predict(test[features])
    mae = mean_absolute_error(test[target], pred)
    mape = mean_absolute_percentage_error(test[target], pred) * 100
    r2 = r2_score(test[target], pred)
    metrics = pd.DataFrame(
        {"Metric": ["MAE (kWh)", "MAPE (%)", "R-squared"], "Value": [round(mae, 3), round(mape, 2), round(r2, 3)]}
    )
    (SUPPORT / "output").mkdir(parents=True, exist_ok=True)
    metrics.to_csv(SUPPORT / "output" / "metrics.csv", index=False)
    print(metrics.to_string(index=False))

    last = test.iloc[-240:]
    fig, ax = plt.subplots(figsize=(9.2, 4.1), facecolor=PAPER)
    ax.plot(last.index, last[target], label="Actual demand", color=INK, linewidth=1.6)
    ax.plot(last.index, pred[-240:], label="Forecast (GBR)", color=PRIMARY, linewidth=1.6, linestyle="--")
    ax.set_xlabel("Timestamp")
    ax.set_ylabel("Load (kWh per half-hour)")
    style_chart(ax, "Figure 4. Actual versus forecast feeder demand (last 5 days of test set)")
    fig.autofmt_xdate()
    fig.tight_layout()
    save(fig, "fig4_forecast_vs_actual.png")

    importances = pd.Series(model.feature_importances_, index=features).sort_values()
    fig, ax = plt.subplots(figsize=(7.2, 4.1), facecolor=PAPER)
    importances.plot(kind="barh", ax=ax, color=PRIMARY, width=0.72)
    ax.set_xlabel("Importance score")
    style_chart(ax, "Figure 5. Relative feature importance — demand forecasting model")
    fig.tight_layout()
    save(fig, "fig5_feature_importance.png")

    profile = df.groupby(["is_weekend", "hour"])["load_kwh"].mean().unstack(level=0)
    fig, ax = plt.subplots(figsize=(8.2, 4.1), facecolor=PAPER)
    ax.plot(profile.index, profile[0], label="Weekday", color=INK, linewidth=1.9)
    ax.plot(profile.index, profile[1], label="Weekend", color=PRIMARY, linewidth=1.9)
    ax.set_xlabel("Hour of day")
    ax.set_ylabel("Average load (kWh)")
    style_chart(ax, "Figure 6. Average half-hourly load profile: weekday vs weekend")
    fig.tight_layout()
    save(fig, "fig6_daily_profile.png")


# ---------------------------------------------------------------------------
# Word report academic fixes
# ---------------------------------------------------------------------------
HEADING1 = {
    2: "Introduction",
    5: "1. Technical, Ethical and Regulatory Requirements",
    25: "2. Evaluation of Cloud Computing Models, Architectures and Processing Approaches",
    38: "3. Designing a Big Data Architecture for MetroEnergy Solutions",
    62: "4. Proof of Concept: Short-Term Demand Forecasting",
    81: "5. Suggestions for Improvement, Governance and Upkeep",
    94: "Conclusion",
    96: "References",
}

HEADING2 = {
    6: "1.1 Data Acquisition Requirements",
    8: "1.2 Storage Requirements",
    10: "1.3 Processing Requirements",
    12: "1.4 Analytical Requirements",
    14: "1.5 Challenges of Big Data for MES",
    19: "1.6 Technical Considerations",
    21: "1.7 Ethical Considerations",
    23: "1.8 Regulatory Considerations",
    26: "2.1 Cloud Deployment Models",
    28: "2.2 Cloud Service Models",
    32: "2.3 Processing Approaches: Batch, Stream, Lambda and Kappa",
    34: "2.4 Data Lake, Data Warehouse and Lakehouse Architectures",
    36: "2.5 Justified Recommendation",
    39: "3.1 Architectural Overview",
    43: "3.2 End-to-End Data Flow",
    48: "3.3 Component Rationale",
    58: "3.4 Technical, Security, Compliance and Performance Responsibilities",
    60: "3.5 Controlling Costs by Design",
    63: "4.1 Objective and Scope",
    65: "4.2 Dataset and Methodology",
    67: "4.3 Implementation Summary",
    70: "4.4 Results",
    79: "4.5 Critical Evaluation of the Proof of Concept",
    82: "5.1 Governance Maturity",
    84: "5.2 Sustainability",
    86: "5.3 Ongoing Maintenance",
    88: "5.4 Innovation Roadmap",
}

CAPTIONS = {
    31: "Figure 1. Shared responsibility split across cloud service models, illustrating why a blended IaaS/PaaS approach suits MES's mixed latency-critical and analytical workloads.",
    42: "Figure 2. Proposed cloud-enabled big data architecture for MetroEnergy Solutions.",
    47: "Figure 3. End-to-end data flow from field devices to business decisions, including a pipeline monitoring feedback loop.",
    74: "Figure 4. Actual versus forecast feeder demand for the last five days of the test set.",
    76: "Figure 5. Relative feature importance in the demand forecasting model; lagged demand and hour-of-day dominate, consistent with the daily periodicity described in Section 1.4.",
    78: "Figure 6. Average half-hourly load profile comparing weekday and weekend patterns, illustrating the seasonal and behavioural structure the model must capture.",
}

MEDIA_MAP = {
    "word/media/image1.png": "fig1_shared_responsibility.png",
    "word/media/image2.png": "fig2_architecture.png",
    "word/media/image3.png": "fig3_dataflow.png",
    "word/media/image4.png": "fig4_forecast_vs_actual.png",
    "word/media/image5.png": "fig5_feature_importance.png",
    "word/media/image6.png": "fig6_daily_profile.png",
}


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def apply_report_fixes():
    backup = SUPPORT / (DOCX.stem + ".BACKUP.docx")
    if not backup.exists():
        shutil.copy2(DOCX, backup)
        print("Backup:", backup.name)

    doc = Document(DOCX)

    if doc.paragraphs[0].text.strip():
        doc.paragraphs[0].style = doc.styles["Title"]
    if doc.paragraphs[1].text.strip():
        try:
            doc.paragraphs[1].style = doc.styles["Subtitle"]
        except KeyError:
            doc.paragraphs[1].style = doc.styles["Normal"]

    for idx, title in HEADING1.items():
        p = doc.paragraphs[idx]
        set_paragraph_text(p, title)
        p.style = doc.styles["Heading 1"]

    for idx, title in HEADING2.items():
        p = doc.paragraphs[idx]
        set_paragraph_text(p, title)
        p.style = doc.styles["Heading 2"]

    try:
        cap_style = doc.styles["Caption"]
    except KeyError:
        cap_style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap_style.font.italic = True
        cap_style.font.size = Pt(10)
        cap_style.font.color.rgb = RGBColor(0x7A, 0x5E, 0x58)

    for idx, cap in CAPTIONS.items():
        p = doc.paragraphs[idx]
        set_paragraph_text(p, cap)
        p.style = cap_style
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x7A, 0x5E, 0x58)

    p40 = doc.paragraphs[40].text
    if "figure 3" in p40.lower():
        set_paragraph_text(
            doc.paragraphs[40],
            p40.replace("figure 3", "Figure 2").replace("Figure 3", "Figure 2"),
        )

    p44 = doc.paragraphs[44].text
    if "Figure 4" in p44:
        set_paragraph_text(doc.paragraphs[44], p44.replace("Figure 4", "Figure 3"))

    tmp = SUPPORT / "_report_text_fixed.docx"
    doc.save(tmp)

    with zipfile.ZipFile(tmp, "r") as zin, zipfile.ZipFile(DOCX, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in MEDIA_MAP:
                src = FIG_DIR / MEDIA_MAP[item.filename]
                data = src.read_bytes()
                print("Replaced", item.filename, "<-", src.name)
            zout.writestr(item, data)
    tmp.unlink(missing_ok=True)
    print("Updated", DOCX.name)


def main():
    print("=== Generating figures ===")
    fig1_shared_responsibility()
    fig2_architecture()
    fig3_dataflow()
    run_poc_and_figs_4_to_6()
    print("=== Fixing Word report ===")
    apply_report_fixes()
    print("Done.")


if __name__ == "__main__":
    main()
